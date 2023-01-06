# -*- coding: utf-8 -*-

from odoo import models, fields, api, tools, http
from odoo.http import request
from odoo.exceptions import Warning

from io import BytesIO
from datetime import date
import base64

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import RGBColor

from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.oxml import ns

from PIL import Image

import logging
_logger = logging.getLogger(__name__)


class PageNumber:
    @staticmethod
    def create_element(name):
        return OxmlElement(name)

    @staticmethod
    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(self, paragraph, start_text):
        page_run = paragraph.add_run()
        t1 = self.create_element('w:t')
        self.create_attribute(t1, 'xml:space', 'preserve')
        t1.text = start_text
        page_run._r.append(t1)

        page_num_run = paragraph.add_run()

        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        of_run = paragraph.add_run()
        t2 = self.create_element('w:t')
        self.create_attribute(t2, 'xml:space', 'preserve')
        t2.text = ' of '
        of_run._r.append(t2)

        fldChar3 = self.create_element('w:fldChar')
        self.create_attribute(fldChar3, 'w:fldCharType', 'begin')

        instrText2 = self.create_element('w:instrText')
        self.create_attribute(instrText2, 'xml:space', 'preserve')
        instrText2.text = "NUMPAGES"

        fldChar4 = self.create_element('w:fldChar')
        self.create_attribute(fldChar4, 'w:fldCharType', 'end')

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)

class CT_Anchor(BaseOxmlElement):
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, r_id, filename, cx, cy, pos_x, pos_y):
        pic_id = 0
        pic = CT_Picture.new(pic_id, filename, r_id, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % (nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y))
        )

class WatermarkPicture():
    def new_pic_anchor(self, part, image_descriptor, width, height, pos_x, pos_y):
        """Return a newly-created `w:anchor` element.
        The element contains the image specified by *image_descriptor* and is scaled
        based on the values of *width* and *height*.
        """
        r_id, image = part.get_or_add_image(image_descriptor)
        cx, cy = image.scaled_dimensions(width, height)
        shape_id, filename = part.next_id, image.filename
        return CT_Anchor.new_pic_anchor(shape_id, r_id, filename, cx, cy, pos_x, pos_y)

    # refer to docx.text.run.add_picture
    def add_watermark_picture(self, p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
        """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
        """
        run = p.add_run()
        anchor = self.new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
        run._r.add_drawing(anchor)

class HrEmployeeResume(models.Model):
    _inherit = 'hr.employee'

    @staticmethod
    def str_date(d):
        if d:
            return d.strftime('%B %d, %Y')
        return ''

    @staticmethod
    def short_str_date(d):
        if d:
            return d.strftime('%B %Y')
        return ''

    @staticmethod
    def image_stream(image_data):
        img = base64.b64decode(image_data)
        image_stream = BytesIO(img)
        return image_stream

    def personal_record(self):
        return {
            'surname': (self.user_partner_id.surname or '') + (self.user_partner_id.surname_prefix or ''),
            'given_name': self.user_partner_id.first_name or '',
            'birthday': self.str_date(self.birthday) or '',
            'nationality': self.country_id.name or ''
        }

    def education_record(self, line):
        return {
            'date_start': line.date_start or '',
            'date_end': line.date_end or '',
            'name': line.name or '',
            'period': self.short_str_date(line.date_start) + ' - ' + self.short_str_date(line.date_end) or '',
            'graduated_on': self.short_str_date(line.date_end) or ''
        }

    def experience_record(self, line):
        return {
            'date_start': line.date_start or '',
            'date_end': line.date_end or '',
            'period': self.short_str_date(line.date_start) + ' - ' + self.short_str_date(line.date_end) or '',
            'company': line.company or '',
            'function': line.name or '',
            'description': line.description or ''
        }

    # level english
    @staticmethod
    def language_record(line):
        return {
            'description': line.description or ''
        }

    def certificate_record(self, line):
        return {
            'date_start': self.short_str_date(line.date_start) or '',
            'date_end': self.short_str_date(line.date_end) or '',
            'description': line.description,
            'name': line.name
        }

    def resume(self):
        result = {
            'personal': self.personal_record(),
            'experience': [],
            'education': [],
            'language': [],
            'certificate': []
        }

        for line in self.resume_line_ids:
            if line.row_display_type == 'study':
                result['education'].append(self.education_record(line))
            if line.row_display_type == 'experience':
                result['experience'].append(self.experience_record(line))
            if line.row_display_type == 'language':
                result['language'].append(self.language_record(line))
            if line.row_display_type == 'cert':
                result['certificate'].append(self.certificate_record(line))

        result['education'].sort(key=lambda k: k['date_start'], reverse=True)
        result['experience'].sort(key=lambda k: k['date_start'], reverse=True)
        return result

    @staticmethod
    def add_text(section, string):
        text = section.add_run(string)
        text.font.size = Pt(12)
        text.font.color.rgb = RGBColor(0, 0, 0)

    def set_row_height(self,row,style):
        # https://stackoverflow.com/questions/37532283/python-how-to-adjust-row-height-of-table-in-docx
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "300")
        trHeight.set(qn('w:hRule'), style)
        trPr.append(trHeight)

    def document_heading(self, section, string, alignment=None):
        heading = section.add_heading()
        self.add_text(heading, string)
        if alignment is not None:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return heading

    @staticmethod
    # alpha = 0-255, lower will result in a more "faded" images, better for watermarks
    def alpha_image(image_stream, alpha):
        img = Image.open(image_stream)
        img.putalpha(alpha)
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        return buffer

    def add_watermark(self, p, image_stream):
        image_stream = self.alpha_image(image_stream, 32)
        register_element_cls('wp:anchor', CT_Anchor)
        watermark = WatermarkPicture()
        for i in range(3):
            watermark.add_watermark_picture(p, image_stream, height=Inches(1.0), pos_x=Pt(200), pos_y=Pt(200 + i*200))

    def add_footer_and_watermark(self, document, watermark_image_stream):
        page_number = PageNumber()
        footer = document.sections[0].footer.paragraphs[0]
        current_date = self.str_date(date.today())
        start_text = 'CV extract, ' + current_date + '\t\tPage '
        page_number.add_page_number(footer, start_text)
        # watermark needs to be added in the special paragraph of the footer,
        # or else it will be applied on the first page of the document
        self.add_watermark(footer, watermark_image_stream)

    @staticmethod
    def add_row_info(table, index, key_text, value_text):
        row = table.rows[index].cells
        row[0].text = key_text
        row[1].text = ''
        row[2].text = value_text

    def add_job_position(self, document):
        table = document.add_table(rows=1, cols=3)
        table.columns[0].width = Mm(40)
        table.columns[1].width = Mm(10)
        table.columns[2].width = Mm(135)
        jp = table.rows[0].cells[0].paragraphs[0].add_run('Job Position')
        jp.font.size = Pt(12)
        jp.bold = True
        table.rows[0].cells[2].paragraphs[0].add_run(self.job_id.name)
        document.add_paragraph('')
        for row in table.rows:
            self.set_row_height(row,"exact")

    def add_personal_information(self, document, resume):
        self.document_heading(document, 'Personal Information')
        # user_image_stream = self.image_stream(self.image_1920)
        # document.add_picture(user_image_stream, width=Inches(1))

        table = document.add_table(rows=4, cols=3)
        table.columns[0].width = Mm(40)
        table.columns[1].width = Mm(10)
        table.columns[2].width = Mm(135)
        self.add_row_info(table, 0, 'Surname', resume['personal']['surname'])
        self.add_row_info(table, 1, 'Given Name', resume['personal']['given_name'])
        self.add_row_info(table, 2, 'Date of Birth', resume['personal']['birthday'])
        self.add_row_info(table, 3, 'Nationality', resume['personal']['nationality'])
        document.add_paragraph('')
        for row in table.rows:
            self.set_row_height(row,"exact")

    def add_work_experience(self, document, resume):
        self.document_heading(document, 'Experience')
        for experience in resume['experience']:
            table = document.add_table(rows=4, cols=3)
            table.columns[0].width = Mm(40)
            table.columns[1].width = Mm(10)
            table.columns[2].width = Mm(135)
            self.add_row_info(table, 0, 'Period', experience['period'])
            self.add_row_info(table, 1, 'Company', experience['company'])
            self.add_row_info(table, 2, 'Function', experience['function'])
            self.add_row_info(table, 3, 'Main Tasks', experience['description'])
            document.add_paragraph('')
            rownum = 0
            for row in table.rows:
                if rownum == 3:
                    self.set_row_height(row,"auto")
                else:
                    self.set_row_height(row, "exact")
                rownum = rownum + 1

    def add_education(self, document, resume):
        self.document_heading(document, 'Education')
        for education in resume['education']:
            table = document.add_table(rows=3, cols=3)
            table.columns[0].width = Mm(40)
            table.columns[1].width = Mm(10)
            table.columns[2].width = Mm(135)
            self.add_row_info(table, 0, 'Description', education['name'])
            self.add_row_info(table, 1, 'Period', education['period'])
            self.add_row_info(table, 2, 'Graduated on date', education['graduated_on'])
            document.add_paragraph('')
            for row in table.rows:
                self.set_row_height(row,"exact")

    def add_language(self, document, resume):
        self.document_heading(document, 'Language')
        for language in resume['language']:
            table = document.add_table(rows=1, cols=3)
            table.columns[0].width = Mm(40)
            table.columns[1].width = Mm(10)
            table.columns[2].width = Mm(135)
            self.add_row_info(table, 0, 'Level english', language['description'])
            for row in table.rows:
                self.set_row_height(row,"exact")

    def add_certificate(self, document, resume):
        self.document_heading(document, 'Certificate')
        for certificate in resume['certificate']:
            table = document.add_table(rows=3, cols=3)
            table.columns[0].width = Mm(40)
            table.columns[1].width = Mm(10)
            table.columns[2].width = Mm(135)
            self.add_row_info(table, 0, 'Certificate', certificate['name'])
            self.add_row_info(table, 1, 'Description', certificate['description'])
            self.add_row_info(table, 2, 'Achieved', str(certificate['date_start']) or 'N/A')
            #self.add_row_info(table, 3, 'Valid Until',str(certificate['date_end']) or 'N/A')
            document.add_paragraph('')
            rownum = 0
            for row in table.rows:
                if rownum == 1:
                    self.set_row_height(row, "auto")
                else:
                    self.set_row_height(row, "exact")
                rownum = rownum + 1

    def resume_document(self, company):
        resume = self.resume()
        document = Document()

        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.25)
            section.bottom_margin = Inches(0.25)
            section.left_margin = Mm(17)
            section.right_margin = Mm(15)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        # logo_image_stream = self.image_stream(request.env.user.company_id.logo)
        logo_image_stream = self.image_stream(company.logo)
        #document.add_picture(logo_image_stream, width=Mm(50))



        section = document.sections[0]
        sec_header = section.header
        header_tp = sec_header.add_paragraph()
        header_run = header_tp.add_run()
        header_run.add_picture(logo_image_stream, width=Inches(2.0))

        self.document_heading(document, 'Curriculum Vitae Extract', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        document.add_paragraph('')

        self.add_job_position(document)
        self.add_personal_information(document, resume)
        self.add_education(document, resume)
        self.add_certificate(document, resume)
        self.add_language(document, resume)
        self.add_work_experience(document, resume)

        self.add_footer_and_watermark(document, logo_image_stream)

        return document

    def data_file(self, company):
        document = self.resume_document(company)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        data = buffer.read()
        data = base64.b64encode(data)
        buffer.close()
        return data

    def generate_resume(self, company):
        # if no company is selected, show a message
        # can be changed to still generate the resume without any logo (or with a generic logo), if needed
        if not company.id:
            raise Warning('Please select a company')

        data = self.data_file(company)
        filename = 'resume.docx'
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'datas': data,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        })

        return {
            'target': 'self',
            'type': 'ir.actions.act_url',
            'url': '/web/content/' + str(attachment.id) + '?download=true'
        }

